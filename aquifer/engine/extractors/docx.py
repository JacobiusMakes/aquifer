"""DOCX text extraction using python-docx."""

from __future__ import annotations

from pathlib import Path


def extract_docx(path: Path) -> str:
    """Extract text from a DOCX file including paragraphs and tables."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install python-docx"
        )

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
